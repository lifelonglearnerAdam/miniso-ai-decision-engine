#!/usr/bin/env python3
"""Build the judge-facing DOCX whitepaper from the versioned Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.font_manager import FontProperties
from matplotlib.patches import FancyBboxPatch

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "02_жҠҖжңҜзҷҪзҡ®д№Ұ" / "AIдә§е“ҒејҖеҸ‘еј•ж“ҺжҠҖжңҜзҷҪзҡ®д№Ұ.md"
OUTPUT = ROOT / "output" / "docx" / "AIдә§е“ҒејҖеҸ‘еј•ж“ҺжҠҖжңҜзҷҪзҡ®д№Ұ.docx"
ASSET_DIR = ROOT / "output" / "qa" / "whitepaper-assets"

# Resolved preset: standard_business_brief.
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120
BODY_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203748"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
BORDER = "D0D5DD"
ACCENT = "D72638"
GOLD = "9A6A00"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_START_END_DXA),
        ("end", CELL_START_END_DXA),
    ):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            set_cell_margins(cell)


def choose_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    lengths = []
    for index in range(columns):
        values = [len(re.sub(r"[*`\[\]]", "", row[index])) for row in rows]
        lengths.append(max(5, min(max(values), 32)))
    minimum = 900 if columns <= 4 else 700
    available = CONTENT_WIDTH_DXA - minimum * columns
    total = sum(lengths)
    widths = [minimum + round(available * value / total) for value in lengths]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def set_run_font(run, size=None, color=None, bold=None, italic=None, font=BODY_FONT) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, bold=False) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    run_props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)
    if bold:
        run_props.append(OxmlElement("w:b"))
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    run_props.append(fonts)
    run_element.append(run_props)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`.+?`|\[[^\]]+\]\(https?://[^)]+\))")


def add_inline(paragraph, text: str, base_size=11, base_color="222222") -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, color=base_color, italic=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size - 0.5, color=DARK_BLUE, font="Consolas")
            run.font.highlight_color = 15
        else:
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=base_size, color=base_color)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(11)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def create_numbering_id(doc: Document) -> int:
    """Create a real single-level decimal list that starts at one."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(element.get(qn("w:numId"))) for element in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    # Word may merge visually identical abstract lists unless they carry
    # distinct namespace/template identifiers, causing later lists to continue
    # the table-of-contents numbering. Keep each Markdown list independent.
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{0xA11C0000 + abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    template = OxmlElement("w:tmpl")
    template.set(qn("w:val"), f"{0xB17E0000 + abstract_id:08X}")
    abstract.append(template)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "decimal")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def format_numbered_paragraph(paragraph, *, space_after: float = 8) -> None:
    """Apply list geometry without Word's built-in List Number style linkage."""
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.167


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=9, color=MUTED)


def populate_header(header) -> None:
    paragraph = header.paragraphs[0]
    paragraph.clear()


def populate_footer(footer) -> None:
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    add_page_number(paragraph)


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    populate_header(section.header)
    populate_footer(section.footer)
    populate_header(section.even_page_header)
    populate_footer(section.even_page_footer)

    first_header = section.first_page_header
    first_header.paragraphs[0].clear()
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].clear()


def create_architecture_diagram(path: Path) -> None:
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    font = FontProperties(fname=str(font_path)) if font_path.exists() else None
    fig, ax = plt.subplots(figsize=(12, 6.2))
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6.2)
    ax.axis("off")
    layers = [
        (5.2, "ж•°жҚ®дёҺзҹҘиҜҶеұӮ", "и¶ӢеҠҝдәӢд»¶ В· е•Ҷе“Ғ/й”ҖйҮҸ В· дҫӣеә”й“ҫ/DFM В· зңҹе®һз ”з©¶ж ·жң¬", "#EAF2F8"),
        (
            4.05,
            "еҶізӯ–жҷәиғҪеұӮ",
            "Trend Agent вҶ’ Creative Agent вҶ’ Evolution + DFM вҶ’ Calibrated Panel",
            "#FDEDEC",
        ),
        (2.90, "иҜ„дј°дёҺеӯҰд№ еұӮ", "ж—¶е…үжңәеӣһжөӢ В· ж ЎеҮҶзӣ‘жҺ§ В· еҒҸеҘҪи®°еҪ• В· Champion/Challenger", "#FEF5E7"),
        (1.75, "еҚҸдҪңдёҺе®Ўжү№еұӮ", "еҖҷйҖүиҜҒжҚ®еҚЎ В· дәәе·Ҙе®Ўжү№ В· з»“жһңеӨҚзӣҳ", "#E9F7EF"),
        (0.60, "жІ»зҗҶе№ійқў", "иә«д»Ҫжқғйҷҗ В· ж•°жҚ®иЎҖзјҳ В· еҶ…е®№е®үе…Ё В· зӣ‘жҺ§йҷҚзә§ В· е®Ўи®Ўеӣһж»ҡ", "#F4F4F5"),
    ]
    for y, title, detail, fill in layers:
        box = FancyBboxPatch(
            (0.55, y),
            10.9,
            0.78,
            boxstyle="round,pad=0.025,rounding_size=0.10",
            linewidth=1.2,
            edgecolor="#B8C0CC",
            facecolor=fill,
        )
        ax.add_patch(box)
        ax.text(
            0.85,
            y + 0.49,
            title,
            fontsize=12.5,
            fontweight="bold",
            color="#203748",
            fontproperties=font,
        )
        ax.text(3.0, y + 0.49, detail, fontsize=10.3, color="#344054", fontproperties=font)
    for y in (5.1, 3.95, 2.8, 1.65):
        ax.annotate(
            "",
            xy=(6, y),
            xytext=(6, y + 0.22),
            arrowprops=dict(arrowstyle="->", color="#D72638", lw=1.4),
        )
    ax.text(
        0.55,
        6.02,
        "дјҒдёҡзә§ AI дә§е“Ғз ”еҸ‘еҶізӯ–жһ¶жһ„",
        fontsize=17,
        fontweight="bold",
        color="#203748",
        fontproperties=font,
    )
    ax.text(
        11.45,
        6.04,
        "жЁЎеһӢеҸҜжӣҝжҚў В· иҜҒжҚ®й“ҫзЁіе®ҡ",
        fontsize=9.5,
        ha="right",
        color="#667085",
        fontproperties=font,
    )
    fig.tight_layout(pad=0.5)
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def add_picture_with_alt(doc: Document, path: Path, alt_text: str, width=6.4) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=Inches(width))._inline
    doc_pr = inline.docPr
    doc_pr.set("descr", alt_text)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    text = caption.add_run("еӣҫ 1  AI дә§е“Ғз ”еҸ‘еҶізӯ–еј•ж“ҺйҖ»иҫ‘жһ¶жһ„")
    set_run_font(text, size=9, color=MUTED, italic=True)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("TECHNICAL WHITE PAPER В· дјҒдёҡй—®йўҳи§ЈеҶіж–№жЎҲиҜ„е®ЎзүҲ")
    set_run_font(run, size=10, color=GOLD, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("AI дә§е“ҒејҖеҸ‘еј•ж“Һ\nжҠҖжңҜзҷҪзҡ®д№Ұ")
    set_run_font(run, size=30, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run("йқўеҗ‘й«ҳйҖҹдёҠж–°зҡ„еҸҜеӣһжөӢгҖҒеҸҜж ЎеҮҶгҖҒдәәеңЁзҺҜдёӯзҡ„ж–°е“Ғз ”еҸ‘еҶізӯ–зі»з»ҹ")
    set_run_font(run, size=15, color=DARK_BLUE)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(34)
    run = rule.add_run("MINISO AI PRODUCT DECISION ENGINE")
    set_run_font(run, size=10.5, color=ACCENT, bold=True)

    metadata = [
        ("зүҲжң¬", "v1.2"),
        ("ж—Ҙжңҹ", "2026 е№ҙ 7 жңҲ 17 ж—Ҙ"),
        ("еҸӮиөӣеӣўйҳҹ", "дёүеҗҚеӯҰз”ҹ"),
        ("йЎ№зӣ®е®ҡдҪҚ", "йқўеҗ‘дјҒдёҡзңҹе®һй—®йўҳзҡ„з«һиөӣйӘҢиҜҒзүҲи§ЈеҶіж–№жЎҲ"),
        ("иҜҒжҚ®еҸЈеҫ„", "е®һзҺ° / еҗҲжҲҗжј”зӨә / дјҒдёҡеҫ…йӘҢиҜҒ"),
        ("д»Јз Ғд»“еә“", "github.com/lifelonglearnerAdam/miniso-ai-decision-engine"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"{label}  В·  ")
        set_run_font(run, size=10, color=MUTED, bold=True)
        run = paragraph.add_run(value)
        set_run_font(run, size=10, color=MUTED)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(28)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("жң¬зҷҪзҡ®д№Ұдёӯзҡ„е®һйӘҢжҢҮж ҮжқҘиҮӘеҗҲжҲҗж•°жҚ®пјҢдёҚд»ЈиЎЁеҗҚеҲӣдјҳе“Ғзңҹе®һдёҡеҠЎ KPIгҖӮ")
    set_run_font(run, size=9.5, color=ACCENT, bold=True)
    doc.add_page_break()


def add_contents(doc: Document, markdown: str) -> None:
    heading = doc.add_paragraph("зӣ®еҪ•", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    sections = [
        match.group(1).strip()
        for match in re.finditer(r"^## (.+)$", markdown, flags=re.MULTILINE)
        if not match.group(1).startswith("йҷ„еҪ•")
    ]
    toc_num_id = create_numbering_id(doc)
    for section in sections:
        paragraph = doc.add_paragraph()
        format_numbered_paragraph(paragraph, space_after=3)
        apply_numbering(paragraph, toc_num_id)
        add_inline(paragraph, re.sub(r"^\d+\.\s*", "", section), base_size=10.5)
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    add_inline(paragraph, "йҷ„еҪ• A-Cпјҡй…ҚзҪ®гҖҒжЈҖжҹҘиЎЁдёҺеҸӮиҖғж–ҮзҢ®", base_size=10.5)

    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(16)
    callout.paragraph_format.space_after = Pt(6)
    callout.paragraph_format.left_indent = Inches(0.18)
    run = callout.add_run("йҳ…иҜ»жҸҗзӨә  ")
    set_run_font(run, size=10.5, color=ACCENT, bold=True)
    run = callout.add_run(
        "е…Ёж–ҮдёҘж јеҢәеҲҶвҖңе·Іе®һзҺ°гҖҒеҗҲжҲҗжј”зӨәгҖҒдјҒдёҡеҫ…йӘҢиҜҒвҖқпјҢд»»дҪ•еҗҲжҲҗз»“жһңеқҮдёҚеҫ—и§ЈйҮҠдёәз”ҹдә§дёҡз»©гҖӮ"
    )
    set_run_font(run, size=10.5, color="344054")
    p_pr = callout._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FFF1F2")
    p_pr.append(shading)
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows or any(len(row) != len(rows[0]) for row in rows):
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = choose_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_index, source_row in enumerate(rows):
        for col_index, value in enumerate(source_row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.10
            add_inline(paragraph, value, base_size=9.2 if len(rows[0]) >= 4 else 9.6)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)


def add_code_block(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shading)
    run = paragraph.add_run(code.strip())
    set_run_font(run, size=8.8, color="344054", font="Consolas")


def add_body_from_markdown(doc: Document, markdown: str, architecture_image: Path) -> None:
    lines = markdown.splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## ж‘ҳиҰҒ")
    index = start
    paragraph_buffer: list[str] = []
    active_numbering_id: int | None = None

    def flush_paragraph(keep_with_next: bool = False) -> None:
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_with_next = keep_with_next
        add_inline(paragraph, text)
        paragraph_buffer.clear()

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped:
            lookahead = index + 1
            while lookahead < len(lines) and not lines[lookahead].strip():
                lookahead += 1
            keep_with_next = lookahead < len(lines) and lines[lookahead].strip().startswith("```")
            flush_paragraph(keep_with_next=keep_with_next)
            active_numbering_id = None
            index += 1
            continue
        if stripped == "---":
            flush_paragraph()
            active_numbering_id = None
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph(keep_with_next=True)
            active_numbering_id = None
            language = stripped[3:].strip()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            if language == "mermaid":
                add_picture_with_alt(
                    doc,
                    architecture_image,
                    "дә”еұӮ AI дә§е“Ғз ”еҸ‘еҶізӯ–жһ¶жһ„пјҡж•°жҚ®гҖҒеҶізӯ–жҷәиғҪгҖҒиҜ„дј°еӯҰд№ гҖҒеҚҸдҪңе®Ўжү№е’ҢжІ»зҗҶе№ійқўгҖӮ",
                )
            else:
                add_code_block(doc, "\n".join(code_lines))
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            active_numbering_id = None
            level = len(heading_match.group(1)) - 1
            text = heading_match.group(2)
            paragraph = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            paragraph.clear()
            # Explicitly reset list geometry. Word can otherwise carry a
            # preceding list's hanging indent into a heading at a page break.
            paragraph.paragraph_format.left_indent = Inches(0)
            paragraph.paragraph_format.right_indent = Inches(0)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            add_inline(
                paragraph,
                text,
                base_size={1: 16, 2: 13, 3: 12}[min(level, 3)],
                base_color=BLUE if level < 3 else DARK_BLUE,
            )
            for run in paragraph.runs:
                run.bold = True
            index += 1
            continue
        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and lines[index + 1].strip().startswith("|")
        ):
            flush_paragraph()
            active_numbering_id = None
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if numbered or bullet:
            flush_paragraph()
            value = (numbered or bullet).group(1)
            value = value.replace("[ ]", "в–Ў").replace("[x]", "вҳ’")
            paragraph = doc.add_paragraph() if numbered else doc.add_paragraph(style="List Bullet")
            if numbered:
                format_numbered_paragraph(paragraph)
                if active_numbering_id is None:
                    active_numbering_id = create_numbering_id(doc)
                apply_numbering(paragraph, active_numbering_id)
            else:
                active_numbering_id = None
            add_inline(paragraph, value)
            index += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            active_numbering_id = None
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(8)
            p_pr = paragraph._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F9FAFB")
            p_pr.append(shading)
            add_inline(paragraph, stripped.lstrip("> "), base_color=DARK_BLUE)
            index += 1
            continue
        active_numbering_id = None
        paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()


def set_document_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "AI дә§е“ҒејҖеҸ‘еј•ж“ҺжҠҖжңҜзҷҪзҡ®д№Ұ"
    props.subject = "дёүдәәеӯҰз”ҹеӣўйҳҹеё®еҠ©дјҒдёҡи§ЈеҶіж–°е“Ғз ”еҸ‘еҶізӯ–й—®йўҳзҡ„з«һиөӣйӘҢиҜҒзүҲж–№жЎҲ"
    props.author = "AI Pioneer Future Talent Competition Three-Student Team"
    props.keywords = "AIдә§е“Ғз ”еҸ‘, еӨҡAgent, ж—¶е…үжңәеӣһжөӢ, дҝқеҪўйў„жөӢ, DFM"
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.settings.odd_and_even_pages_header_footer = True


def build() -> Path:
    markdown = SOURCE.read_text(encoding="utf-8")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture_image = ASSET_DIR / "architecture.png"
    create_architecture_diagram(architecture_image)

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    set_document_properties(doc)
    add_cover(doc)
    add_contents(doc, markdown)
    add_body_from_markdown(doc, markdown, architecture_image)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(path)
