#!/usr/bin/env python3
from __future__ import annotations

import csv
import html
import json
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
WORK = ROOT / "work" / "report_assets"
MD_PATH = OUT / "miniso_user_insight_report_zh.md"
DOCX_PATH = OUT / "miniso_user_insight_report_zh.docx"
PDF_PATH = OUT / "miniso_user_insight_report_zh.pdf"
SUMMARY_PATH = OUT / "miniso_voc_summary.json"
DATA_PATH = OUT / "miniso_voc_dataset.csv"

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
INK = "243447"
MUTED = "65758B"
LIGHT = "F4F6F9"
ACCENT = "E86A4A"


def parse_markdown(path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("title", line[2:].strip()))
            i += 1
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
        elif line.startswith("> "):
            blocks.append(("quote", line[2:].strip()))
            i += 1
        elif line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for idx, raw in enumerate(table_lines):
                cells = [c.strip() for c in raw.strip().strip("|").split("|")]
                if idx == 1 and all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                    continue
                rows.append(cells)
            blocks.append(("table", rows))
        elif re.match(r"^\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i]).strip())
                i += 1
            blocks.append(("numbered", items))
        elif line.startswith("- "):
            items = []
            while i < len(lines) and lines[i].startswith("- "):
                items.append(lines[i][2:].strip())
                i += 1
            blocks.append(("bullets", items))
        else:
            para = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not re.match(
                r"^(#|>|\| |- |\d+\.\s+)", lines[i]
            ):
                para.append(lines[i].strip())
                i += 1
            blocks.append(("p", " ".join(para)))
    return blocks


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_docx(paragraph, text):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def md_inline_to_reportlab(text):
    text = html.escape(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`(.+?)`", r'<font color="#1F4D78">\1</font>', text)
    text = re.sub(
        r"\[([^\]]+)\]\(([^)]+)\)",
        r'<link href="\2" color="#2E74B5"><u>\1</u></link>',
        text,
    )
    return text


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, total_width=9360, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("名创优品用户洞察  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_docx_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in [
        ("Heading 1", 17, BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, DEEP_BLUE, 8, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    quote = doc.styles.add_style("Insight Quote", WD_STYLE_TYPE.PARAGRAPH)
    quote.base_style = normal
    quote.font.size = Pt(11)
    quote.font.bold = True
    quote.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    quote.paragraph_format.left_indent = Inches(0.2)
    quote.paragraph_format.right_indent = Inches(0.2)
    quote.paragraph_format.space_before = Pt(8)
    quote.paragraph_format.space_after = Pt(12)
    quote.paragraph_format.line_spacing = 1.25
    quote.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = quote.element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), ACCENT)
    borders.append(left)
    p_pr.append(borders)


def make_charts(summary):
    WORK.mkdir(parents=True, exist_ok=True)
    font_path = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
    font = ImageFont.truetype(font_path, 28)
    small = ImageFont.truetype(font_path, 22)
    title_font = ImageFont.truetype(font_path, 34)

    def bar_chart(items, title, path, color):
        width, height = 1500, 760
        image = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(image)
        draw.text((80, 45), title, font=title_font, fill="#243447")
        max_value = max(v for _, v in items)
        y = 125
        for label, value in items:
            draw.text((80, y + 6), label, font=font, fill="#243447")
            bar_x = 390
            bar_w = int(900 * value / max_value)
            draw.rounded_rectangle((bar_x, y, bar_x + bar_w, y + 44), radius=6, fill=color)
            draw.text((bar_x + bar_w + 18, y + 5), str(value), font=small, fill="#65758B")
            y += 72
        draw.text((80, height - 45), "注：为本次目的性公开样本数量，不代表总体市场占比。", font=small, fill="#65758B")
        image.save(path, quality=95)

    platform = sorted(summary["by_platform"].items(), key=lambda x: x[1], reverse=True)
    theme = sorted(summary["by_theme"].items(), key=lambda x: x[1], reverse=True)[:8]
    p1 = WORK / "platform_distribution.png"
    p2 = WORK / "theme_distribution.png"
    bar_chart(platform, "进入分析样本的平台分布", p1, "#2E74B5")
    bar_chart(theme, "高频用户反馈主题（Top 8）", p2, "#E86A4A")
    return p1, p2


def build_docx(blocks, chart_paths):
    doc = Document()
    setup_docx_styles(doc)
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    add_page_number(section.footer.paragraphs[0])
    header_p = section.header.paragraphs[0]
    header_p.text = "公开用户之声（VoC）与 AI 产品开发决策应用"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    title = blocks[0][1]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Arial Unicode MS"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(18)
    add_inline_docx(p2, "中国社交媒体增强版 | 公开用户之声（VoC）与 AI 产品开发决策应用")
    for run in p2.runs:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string(BLUE)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(110)
    p3.add_run("飞书 AI 比赛专题研究\n采集日期：2026-07-18\n样本：124 条公开反馈 | 8 类平台来源").font.size = Pt(11)
    doc.add_page_break()

    # Compact, manually curated table of contents for predictable rendering.
    toc = doc.add_heading("目录", level=1)
    for text in [
        "1. 执行摘要", "2. 研究范围与证据边界", "3. 样本概览", "4. 中国市场核心洞察",
        "5. 海外反馈的补充价值", "6. 代表性证据卡", "7. 用户任务与人群", "8. 产品机会优先级",
        "9. AI 产品开发智能决策引擎", "10. 飞书 AI 落地设计", "11. 比赛展示建议", "12. 数据文件与结论",
    ]:
        q = doc.add_paragraph(style="List Bullet")
        q.paragraph_format.left_indent = Inches(0.25)
        q.paragraph_format.space_after = Pt(4)
        q.add_run(text)
    doc.add_page_break()

    chart_inserted = False
    for kind, payload in blocks[1:]:
        if kind == "title":
            continue
        if kind == "h2":
            doc.add_heading(payload, level=1)
        elif kind == "h3":
            doc.add_heading(payload, level=2)
        elif kind == "quote":
            q = doc.add_paragraph(style="Insight Quote")
            add_inline_docx(q, payload)
        elif kind == "p":
            p = doc.add_paragraph()
            add_inline_docx(p, payload)
            if payload.startswith("进入分析的 119 条样本") and not chart_inserted:
                for chart in chart_paths:
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cp.paragraph_format.space_after = Pt(8)
                    cp.add_run().add_picture(str(chart), width=Inches(6.15))
                chart_inserted = True
        elif kind in ("bullets", "numbered"):
            style = "List Bullet" if kind == "bullets" else "List Number"
            for item in payload:
                p = doc.add_paragraph(style=style)
                p.paragraph_format.left_indent = Inches(0.34)
                p.paragraph_format.first_line_indent = Inches(-0.17)
                p.paragraph_format.space_after = Pt(4)
                add_inline_docx(p, item)
        elif kind == "table":
            rows = payload
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            if len(rows[0]) == 3:
                widths = [2520, 3420, 3420]
            elif len(rows[0]) == 4:
                widths = [1050, 1900, 3010, 3400]
            elif len(rows[0]) == 5:
                widths = [1050, 900, 1700, 2700, 3010]
            else:
                base = 9360 // len(rows[0])
                widths = [base] * (len(rows[0]) - 1) + [9360 - base * (len(rows[0]) - 1)]
            set_table_geometry(table, widths)
            for r_idx, row in enumerate(rows):
                for c_idx, value in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.text = ""
                    add_inline_docx(cell.paragraphs[0], value)
                    set_cell_margins(cell)
                    if r_idx == 0:
                        set_cell_shading(cell, DEEP_BLUE)
                        for run in cell.paragraphs[0].runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    elif r_idx % 2 == 0:
                        set_cell_shading(cell, LIGHT)
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(2)
                        paragraph.paragraph_format.line_spacing = 1.08
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.size = Pt(8.3)
            doc.add_paragraph().paragraph_format.space_after = Pt(1)

    props = doc.core_properties
    props.title = title
    props.subject = "名创优品公开用户反馈与 AI 产品开发决策"
    props.author = "Codex 辅助研究"
    props.keywords = "名创优品, MINISO, 用户洞察, VoC, 小红书, 飞书AI"
    doc.save(DOCX_PATH)


def pdf_page(canvas, doc):
    canvas.saveState()
    canvas.setFont("ArialUnicode", 8)
    canvas.setFillColor(colors.HexColor("#65758B"))
    canvas.drawString(18 * mm, 12 * mm, "名创优品用户真实反馈与产品机会研究")
    canvas.drawRightString(192 * mm, 12 * mm, str(doc.page))
    canvas.setStrokeColor(colors.HexColor("#D9E1EA"))
    canvas.line(18 * mm, 17 * mm, 192 * mm, 17 * mm)
    canvas.restoreState()


def build_pdf(blocks, chart_paths):
    registerFont(TTFont("ArialUnicode", "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"))
    page_w, page_h = A4
    frame = Frame(18 * mm, 20 * mm, page_w - 36 * mm, page_h - 36 * mm, id="normal")
    template = PageTemplate(id="content", frames=frame, onPage=pdf_page)
    doc = BaseDocTemplate(
        str(PDF_PATH), pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=20 * mm, title=blocks[0][1], author="Codex 辅助研究"
    )
    doc.addPageTemplates([template])
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CNBody", parent=styles["BodyText"], fontName="ArialUnicode", fontSize=9.5,
        leading=15, textColor=colors.HexColor("#243447"), alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    h1 = ParagraphStyle(
        "CNH1", parent=body, fontSize=17, leading=22, textColor=colors.HexColor("#2E74B5"),
        spaceBefore=12, spaceAfter=7, keepWithNext=True,
    )
    h2 = ParagraphStyle(
        "CNH2", parent=body, fontSize=12.5, leading=17, textColor=colors.HexColor("#1F4D78"),
        spaceBefore=8, spaceAfter=5, keepWithNext=True,
    )
    quote = ParagraphStyle(
        "CNQuote", parent=body, fontSize=10.5, leading=16, leftIndent=8 * mm, rightIndent=6 * mm,
        borderColor=colors.HexColor("#E86A4A"), borderWidth=0, borderLeft=3,
        borderPadding=7, textColor=colors.HexColor("#1F4D78"), spaceBefore=8, spaceAfter=10,
    )
    cover_title = ParagraphStyle(
        "CoverTitle", parent=body, fontSize=26, leading=36, alignment=TA_CENTER,
        textColor=colors.HexColor("#1F4D78"), spaceAfter=18,
    )
    cover_sub = ParagraphStyle(
        "CoverSub", parent=body, fontSize=13, leading=21, alignment=TA_CENTER,
        textColor=colors.HexColor("#2E74B5"),
    )
    small = ParagraphStyle("Small", parent=body, fontSize=7.6, leading=10.5, alignment=TA_LEFT)

    story = [
        Spacer(1, 52 * mm),
        Paragraph(md_inline_to_reportlab(blocks[0][1]), cover_title),
        Paragraph("中国社交媒体增强版 | 公开用户之声（VoC）与 AI 产品开发决策应用", cover_sub),
        Spacer(1, 42 * mm),
        Paragraph("飞书 AI 比赛专题研究<br/>采集日期：2026-07-18<br/>样本：124 条公开反馈 | 8 类平台来源", cover_sub),
        PageBreak(),
        Paragraph("目录", h1),
    ]
    toc_items = [
        "1. 执行摘要", "2. 研究范围与证据边界", "3. 样本概览", "4. 中国市场核心洞察",
        "5. 海外反馈的补充价值", "6. 代表性证据卡", "7. 用户任务与人群", "8. 产品机会优先级",
        "9. AI 产品开发智能决策引擎", "10. 飞书 AI 落地设计", "11. 比赛展示建议", "12. 数据文件与结论",
    ]
    story.append(ListFlowable([ListItem(Paragraph(x, body), leftIndent=5 * mm) for x in toc_items], bulletType="bullet"))
    story.append(PageBreak())

    chart_inserted = False
    for kind, payload in blocks[1:]:
        if kind == "h2":
            story.append(Paragraph(md_inline_to_reportlab(payload), h1))
        elif kind == "h3":
            story.append(Paragraph(md_inline_to_reportlab(payload), h2))
        elif kind == "quote":
            story.append(Paragraph(md_inline_to_reportlab(payload), quote))
        elif kind == "p":
            story.append(Paragraph(md_inline_to_reportlab(payload), body))
            if payload.startswith("进入分析的 119 条样本") and not chart_inserted:
                for chart in chart_paths:
                    story.append(RLImage(str(chart), width=169 * mm, height=85.6 * mm))
                    story.append(Spacer(1, 3 * mm))
                chart_inserted = True
        elif kind in ("bullets", "numbered"):
            items = [ListItem(Paragraph(md_inline_to_reportlab(x), body), leftIndent=5 * mm) for x in payload]
            story.append(ListFlowable(items, bulletType="1" if kind == "numbered" else "bullet", leftIndent=5 * mm))
            story.append(Spacer(1, 2 * mm))
        elif kind == "table":
            rows = [[Paragraph(md_inline_to_reportlab(c), small) for c in row] for row in payload]
            cols = len(rows[0])
            avail = 174 * mm
            if cols == 3:
                widths = [avail * 0.27, avail * 0.37, avail * 0.36]
            elif cols == 4:
                widths = [avail * 0.14, avail * 0.12, avail * 0.34, avail * 0.40]
            elif cols == 5:
                widths = [avail * 0.11, avail * 0.10, avail * 0.18, avail * 0.31, avail * 0.30]
            else:
                widths = [avail / cols] * cols
            table = Table(rows, colWidths=widths, repeatRows=1, hAlign="CENTER")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, -1), "ArialUnicode"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D3DF")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F4F6F9")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(table)
            story.append(Spacer(1, 3 * mm))

    doc.build(story)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    summary = json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
    blocks = parse_markdown(MD_PATH)
    charts = make_charts(summary)
    build_docx(blocks, charts)
    build_pdf(blocks, charts)
    print(f"Created {DOCX_PATH}")
    print(f"Created {PDF_PATH}")


if __name__ == "__main__":
    main()
